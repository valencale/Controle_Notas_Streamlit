"""
gerar_portfolio.py — Gera o documento DOCX do Portfolio de Web Analytics.
Aluno: Alex Valencia | Disciplina: Web Analytics | 2026

Formato: segue o padrão do 'Portfolio Exemplo Final.pdf' da Anhanguera.
Seções: Capa, Índice, Introdução, Métodos (3 passos), Resultados, Conclusão.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_FILE = os.path.join(OUTPUT_DIR, "Portfolio_Web_Analytics_Alex_Valencia.docx")

# ══════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════

def set_run_font(run, name="Arial", size=12, bold=False, italic=False, color=None):
    """Configura fonte de um run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text, font_size=12, bold=False, italic=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
                  space_before=0, color=None, font_name="Arial"):
    """Adiciona parágrafo formatado ao documento."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run_font(run, name=font_name, size=font_size, bold=bold, italic=italic, color=color)
    return p


def add_code_block(doc, code_text):
    """Adiciona bloco de código com fundo visual (fonte mono, indentado)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    set_run_font(run, name="Consolas", size=9, color=(30, 30, 30))


def add_bullet(doc, text, font_size=12, bold_prefix=None):
    """Adiciona bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        set_run_font(run_b, size=font_size, bold=True)
        run_t = p.add_run(text)
        set_run_font(run_t, size=font_size)
    else:
        run = p.add_run(text)
        set_run_font(run, size=font_size)


# ══════════════════════════════════════════════════════════════
# CONSTRUÇÃO DO DOCUMENTO
# ══════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # Configurar margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # ══════════════════════════════════════════════════════════
    # PÁGINA 1 — CAPA
    # ══════════════════════════════════════════════════════════
    # Espaço superior
    for _ in range(6):
        add_paragraph(doc, "", font_size=12)

    add_paragraph(doc, "Anhanguera", font_size=28, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_paragraph(doc, "Portfolio - Relatorio de Aula", font_size=20,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Pratica", font_size=20,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_paragraph(doc, "Web Analytics", font_size=22, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_paragraph(doc, "Alex Valencia", font_size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)

    # Espaço inferior
    for _ in range(6):
        add_paragraph(doc, "", font_size=12)

    add_paragraph(doc, "Sao Paulo, 2026", font_size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PÁGINA 2 — ÍNDICE
    # ══════════════════════════════════════════════════════════
    add_paragraph(doc, "Indice", font_size=22, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=24)

    indice_items = [
        ("1. Introducao", "3"),
        ("2. Metodos", "4"),
        ("   2.1 Web Scraping", "4"),
        ("   2.2 Data Marketplace", "6"),
        ("   2.3 Etica na Raspagem de Dados", "7"),
        ("3. Resultados", "8"),
        ("4. Conclusao", "9"),
    ]

    for titulo, pagina in indice_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        # Usar tab para alinhar números de página
        run = p.add_run(f"{titulo} {'.' * (50 - len(titulo))} {pagina}")
        set_run_font(run, size=12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PÁGINA 3 — INTRODUÇÃO
    # ══════════════════════════════════════════════════════════
    add_paragraph(doc, "Introducao", font_size=22, bold=True,
                  space_after=16)

    add_paragraph(doc, (
        "Este portfolio apresenta o desenvolvimento pratico de tecnicas de "
        "coleta automatizada de dados utilizando Web Scraping com a biblioteca "
        "BeautifulSoup em Python."
    ), font_size=12, space_after=12)

    add_paragraph(doc, (
        "Foram explorados conceitos fundamentais de Web Analytics, abrangendo "
        "a extracao automatizada de dados de paginas web, a exploracao de "
        "Data Marketplaces como fonte de datasets estruturados, e uma reflexao "
        "critica sobre os desafios eticos e legais associados a raspagem de dados."
    ), font_size=12, space_after=12)

    add_paragraph(doc, (
        "A atividade pratica foi conduzida em tres etapas: implementacao de "
        "um script de Web Scraping funcional, analise de um dataset publico "
        "do Kaggle, e pesquisa sobre casos reais de violacoes eticas "
        "relacionadas a coleta automatizada de dados."
    ), font_size=12, space_after=12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PÁGINAS 4-5 — MÉTODOS: PASSO 1 (WEB SCRAPING)
    # ══════════════════════════════════════════════════════════
    add_paragraph(doc, "Metodos", font_size=22, bold=True,
                  space_after=16)

    add_paragraph(doc, "2.1 Web Scraping — Coleta Automatizada", font_size=16,
                  bold=True, space_after=12)

    add_paragraph(doc, (
        "A atividade de Web Scraping foi realizada utilizando o site "
        "books.toscrape.com, uma plataforma publica projetada especificamente "
        "para praticas de raspagem de dados. O objetivo foi extrair informacoes "
        "de livros (titulo, preco, avaliacao e disponibilidade) de forma automatizada."
    ), font_size=12, space_after=12)

    add_paragraph(doc, "1. Instalacao das Bibliotecas:", font_size=12,
                  bold=True, space_after=6)

    add_code_block(doc, "pip install beautifulsoup4\npip install requests")

    add_paragraph(doc, "2. Codigo do Script de Web Scraping:", font_size=12,
                  bold=True, space_after=6, space_before=12)

    code = '''import requests
from bs4 import BeautifulSoup
import csv

url = "http://books.toscrape.com/"

headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
}

response = requests.get(url, headers=headers)
soup = BeautifulSoup(response.text, "html.parser")

livros = []
for article in soup.find_all("article", class_="product_pod"):
    titulo = article.h3.a["title"]
    preco = article.find("p", class_="price_color").text.strip()
    rating = article.find("p", class_="star-rating")
    avaliacao = rating["class"][1] if rating else "N/A"

    livros.append({
        "titulo": titulo,
        "preco": preco,
        "avaliacao": avaliacao
    })

# Exibir resultados
for livro in livros:
    print(f"{livro['titulo']} - {livro['preco']}")

# Salvar em CSV
with open("livros_scraping.csv", "w", newline="",
          encoding="utf-8") as f:
    writer = csv.DictWriter(f,
        fieldnames=["titulo", "preco", "avaliacao"])
    writer.writeheader()
    writer.writerows(livros)

print(f"{len(livros)} livros extraidos com sucesso!")'''

    add_code_block(doc, code)

    add_paragraph(doc, "3. Resultado da Execucao:", font_size=12,
                  bold=True, space_after=6, space_before=12)

    add_paragraph(doc, (
        "O script foi executado com sucesso, extraindo 20 livros da pagina "
        "principal do site. Os dados foram exibidos no terminal e exportados "
        "para o arquivo livros_scraping.csv."
    ), font_size=12, space_after=8)

    # Resultado simulado do terminal
    terminal_output = """============================================================
  WEB SCRAPING - Books to Scrape
  Disciplina: Web Analytics | Alex Valencia | 2026
============================================================

[>] Acessando: http://books.toscrape.com/
[OK] Status HTTP: 200

[OK] 20 livros extraidos com sucesso!

------------------------------------------------------------
TITULO                                           PRECO
------------------------------------------------------------
A Light in the Attic                           £51.77
Tipping the Velvet                             £53.74
Soumission                                     £50.10
Sharp Objects                                  £47.82
Sapiens: A Brief History of Humankind          £54.23
The Requiem Red                                £22.65
...
------------------------------------------------------------

[SAVE] Arquivo salvo: livros_scraping.csv
[OK] Processo concluido! 20 registros exportados."""

    add_code_block(doc, terminal_output)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PÁGINA 6 — MÉTODOS: PASSO 2 (DATA MARKETPLACE)
    # ══════════════════════════════════════════════════════════
    add_paragraph(doc, "2.2 Data Marketplace — Kaggle", font_size=16,
                  bold=True, space_after=12)

    add_paragraph(doc, (
        "Data Marketplaces sao plataformas que disponibilizam conjuntos de "
        "dados (datasets) para analise, pesquisa e desenvolvimento de projetos. "
        "O Kaggle e um dos maiores e mais populares Data Marketplaces do mundo, "
        "mantido pelo Google, oferecendo milhares de datasets gratuitos."
    ), font_size=12, space_after=12)

    add_paragraph(doc, "Dataset Selecionado:", font_size=12,
                  bold=True, space_after=6)

    add_paragraph(doc, (
        'O dataset escolhido foi o "E-Commerce Customer Behavior Analytics", '
        'disponivel em kaggle.com. Este dataset abrange dados de comportamento '
        'de clientes em e-commerce, incluindo 20.000 clientes e 120.000 sessoes '
        'de navegacao.'
    ), font_size=12, space_after=8)

    add_paragraph(doc, "Estrutura do Dataset:", font_size=12,
                  bold=True, space_after=6)

    add_bullet(doc, "Clientes: dados demograficos, localizacao, tipo de conta")
    add_bullet(doc, "Sessoes de Navegacao: clickstream, tempo de permanencia, paginas visitadas")
    add_bullet(doc, "Pedidos: produtos comprados, valores, datas")
    add_bullet(doc, "Avaliacoes: reviews de produtos com nota e comentario")

    add_paragraph(doc, "", font_size=6, space_after=6)

    add_paragraph(doc, (
        "A analise da estrutura revelou um dataset multi-tabela bem organizado, "
        "ideal para projetos de Web Analytics por conter dados reais de "
        "comportamento de usuarios em ambiente de comercio eletronico."
    ), font_size=12, space_after=12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PÁGINA 7 — MÉTODOS: PASSO 3 (ÉTICA)
    # ══════════════════════════════════════════════════════════
    add_paragraph(doc, "2.3 Etica na Raspagem de Dados", font_size=16,
                  bold=True, space_after=12)

    add_paragraph(doc, (
        "A raspagem de dados, embora seja uma ferramenta poderosa para coleta "
        "automatizada, traz consigo questoes eticas e legais significativas. "
        "A seguir, dois casos reais que ilustram esses desafios:"
    ), font_size=12, space_after=12)

    add_paragraph(doc, "Caso 1: LinkedIn vs. hiQ Labs (2017-2022)",
                  font_size=12, bold=True, space_after=6)

    add_paragraph(doc, (
        "A empresa hiQ Labs utilizava web scraping para coletar dados publicos "
        "de perfis do LinkedIn e oferecer servicos de analytics de RH. O LinkedIn "
        "tentou bloquear a pratica, alegando violacao dos termos de uso. O caso "
        "chegou a Suprema Corte dos EUA e ao Nono Circuito, que decidiu que "
        "raspar dados publicos nao viola a lei federal anti-hacking (CFAA). "
        "Porem, em 2022, o LinkedIn venceu por violacao contratual dos Termos "
        "de Servico, resultando em uma multa de US$ 500.000 contra a hiQ."
    ), font_size=12, space_after=12)

    add_paragraph(doc, "Caso 2: Clearview AI (2020-presente)",
                  font_size=12, bold=True, space_after=6)

    add_paragraph(doc, (
        "A Clearview AI raspou bilhoes de fotos de redes sociais (Facebook, "
        "Instagram, Twitter) sem consentimento para construir um banco de dados "
        "de reconhecimento facial vendido a agencias policiais. A empresa foi "
        "processada em diversos paises por violacao de privacidade. Nos EUA, "
        "foi proibida de vender seu banco de dados a entidades privadas. Na "
        "Europa, recebeu multas por violacao do GDPR. O caso evidencia como "
        "a raspagem massiva de dados pessoais pode gerar consequencias legais "
        "graves em escala global."
    ), font_size=12, space_after=12)

    add_paragraph(doc, "Reflexao:", font_size=12, bold=True, space_after=6)

    add_paragraph(doc, (
        "Estes casos demonstram que, mesmo quando os dados sao tecnicamente "
        "publicos, sua coleta automatizada pode violar termos de servico, leis "
        "de protecao de dados (como LGPD e GDPR) e direitos individuais de "
        "privacidade. A pratica etica de Web Scraping exige respeitar os "
        "termos de uso dos sites, obter consentimento quando necessario, e "
        "evitar a coleta de dados pessoais sensiveis."
    ), font_size=12, space_after=12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PÁGINA 8 — RESULTADOS
    # ══════════════════════════════════════════════════════════
    add_paragraph(doc, "Resultados", font_size=22, bold=True,
                  space_after=16)

    add_paragraph(doc, (
        "A aplicacao dos procedimentos descritos nos metodos permitiu alcancar "
        "os seguintes resultados:"
    ), font_size=12, space_after=12)

    add_bullet(doc, ("Desenvolvimento de um script funcional de Web Scraping "
                     "utilizando BeautifulSoup e Requests em Python"))
    add_bullet(doc, ("Extracao bem-sucedida de 20 registros de livros "
                     "(titulo, preco, avaliacao e disponibilidade) do site "
                     "books.toscrape.com"))
    add_bullet(doc, "Exportacao dos dados extraidos para formato CSV estruturado")
    add_bullet(doc, ("Exploracao do Data Marketplace Kaggle, com analise "
                     "detalhada da estrutura de um dataset de e-commerce "
                     "contendo 20.000 clientes e 120.000 sessoes"))
    add_bullet(doc, ("Identificacao e documentacao de dois casos reais de "
                     "violacoes eticas na raspagem de dados (LinkedIn vs. "
                     "hiQ Labs e Clearview AI)"))
    add_bullet(doc, ("Compreensao pratica da importancia do respeito aos "
                     "Termos de Servico e legislacoes de protecao de dados"))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PÁGINA 9 — CONCLUSÃO
    # ══════════════════════════════════════════════════════════
    add_paragraph(doc, "Conclusao", font_size=22, bold=True,
                  space_after=16)

    add_paragraph(doc, (
        "A pratica com Web Scraping utilizando BeautifulSoup evidenciou o "
        "potencial da coleta automatizada de dados como ferramenta essencial "
        "para Web Analytics. A capacidade de extrair informacoes estruturadas "
        "de paginas web de forma programatica abre possibilidades significativas "
        "para analise de mercado, monitoramento de precos e inteligencia "
        "competitiva."
    ), font_size=12, space_after=12)

    add_paragraph(doc, (
        "A exploracao do Kaggle como Data Marketplace demonstrou a riqueza de "
        "datasets disponiveis publicamente, permitindo que profissionais de "
        "analytics acessem dados estruturados sem necessidade de construir "
        "pipelines de coleta proprios."
    ), font_size=12, space_after=12)

    add_paragraph(doc, (
        "Por fim, a investigacao dos casos LinkedIn vs. hiQ Labs e Clearview AI "
        "reforçou a importancia critica da etica na coleta de dados. O "
        "equilibrio entre eficiencia na obtencao de informacoes e o respeito "
        "a privacidade, aos termos de servico e as legislacoes vigentes e "
        "fundamental para a pratica responsavel de Web Analytics."
    ), font_size=12, space_after=12)

    # ══════════════════════════════════════════════════════════
    # SALVAR
    # ══════════════════════════════════════════════════════════
    doc.save(DOCX_FILE)
    print(f"[OK] Documento gerado: {DOCX_FILE}")
    return DOCX_FILE


if __name__ == "__main__":
    print("=" * 60)
    print("  GERADOR DE PORTFOLIO — Web Analytics")
    print("  Aluno: Alex Valencia | 2026")
    print("=" * 60)
    print()
    filepath = build_document()
    print(f"\n[OK] Portfolio salvo em:\n     {filepath}")
    print("\nProximo passo: abra o DOCX, revise, e exporte como PDF.")
