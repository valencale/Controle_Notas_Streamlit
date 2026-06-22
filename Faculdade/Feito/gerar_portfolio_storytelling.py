"""
gerar_portfolio_storytelling.py — Gera DOCX do Portfolio de Visualizacao de Dados.
Aluno: Alex Valencia | 2026
"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DIR = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(DIR, "Portfolio_Visualizacao_Dados_Alex_Valencia.docx")

def rf(run, name="Arial", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def ap(doc, text, sz=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=6, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    rf(r, size=sz, bold=bold)
    return p

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    rf(r, name="Consolas", size=9)

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    rf(r, size=12)

def img(doc, filename, width=Inches(5.5)):
    path = os.path.join(DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2)

    # ===== CAPA =====
    for _ in range(6):
        ap(doc, "", sz=12)
    ap(doc, "Anhanguera", sz=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    ap(doc, "Portfolio - Relatorio de Aula", sz=20, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    ap(doc, "Pratica", sz=20, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    ap(doc, "Visualizacao de Dados e Storytelling", sz=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    ap(doc, "Alex Valencia", sz=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=48)
    for _ in range(6):
        ap(doc, "", sz=12)
    ap(doc, "Sao Paulo, 2026", sz=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()

    # ===== INDICE =====
    ap(doc, "Indice", sz=22, bold=True, after=24)
    for t, p in [("1. Introducao","3"),("2. Metodos","4"),
                 ("   2.1 Carregamento dos Dados","4"),
                 ("   2.2 Grafico de Linha - Tendencias Anuais","5"),
                 ("   2.3 Grafico de Barras - Consumo Mensal","6"),
                 ("   2.4 Analise de Correlacao","7"),
                 ("   2.5 Grafico Interativo com Plotly","8"),
                 ("3. Resultados","9"),("4. Conclusao","10")]:
        ap(doc, f"{t} {'.' * (50 - len(t))} {p}", sz=12, after=8)
    doc.add_page_break()

    # ===== INTRODUCAO =====
    ap(doc, "Introducao", sz=22, bold=True, after=16)
    ap(doc, (
        "Este portfolio apresenta o desenvolvimento pratico de tecnicas de "
        "visualizacao de dados e storytelling aplicadas a analise do consumo "
        "de energia eletrica do Palacio do Jaburu, no periodo de 2017 a 2024."
    ), after=12)
    ap(doc, (
        "Foram exploradas diferentes formas de visualizacao grafica utilizando "
        "as bibliotecas Matplotlib e Plotly em Python, com o objetivo de "
        "identificar tendencias anuais, padroes sazonais e correlacoes entre "
        "o consumo energetico e o valor das faturas."
    ), after=12)
    ap(doc, (
        "A atividade foi conduzida seguindo os procedimentos do roteiro da "
        "disciplina, utilizando o Google Colab como ambiente de desenvolvimento "
        "e o dataset publico disponibilizado pelo professor."
    ), after=12)
    doc.add_page_break()

    # ===== METODOS =====
    ap(doc, "Metodos", sz=22, bold=True, after=16)

    # 2.1 Carregamento
    ap(doc, "2.1 Carregamento dos Dados", sz=16, bold=True, after=12)
    ap(doc, (
        "O conjunto de dados foi obtido do repositorio GitHub do professor, "
        "contendo 88 registros de consumo mensal de energia eletrica do "
        "Palacio do Jaburu entre junho de 2017 e setembro de 2024."
    ), after=8)
    ap(doc, "Estrutura do Dataset:", sz=12, bold=True, after=6)
    bullet(doc, "ANO: ano de referencia (2017 a 2024)")
    bullet(doc, "MES: mes de referencia")
    bullet(doc, "CONSUMO-KWH: consumo em quilowatt-hora")
    bullet(doc, "VALOR BRUTO DA FATURA: valor monetario em reais (R$)")

    ap(doc, "Codigo de importacao e limpeza:", sz=12, bold=True, after=6, before=12)
    code(doc, """import pandas as pd
import matplotlib.pyplot as plt
import plotly.express as px

# Carregar dados
df = pd.read_csv('consumo_energia_jaburu.csv',
                 sep=';', encoding='latin-1')
df.columns = ['ANO','MES','CONSUMO_KWH','VALOR_FATURA']

# Converter valor: "R$ 17.910,77" -> float
df['VALOR'] = (df['VALOR_FATURA']
    .str.replace('R$','',regex=False)
    .str.replace('.','',regex=False)
    .str.replace(',','.',regex=False)
    .str.strip().astype(float))

print(df.shape)  # (88, 5)
print(df.head())""")
    doc.add_page_break()

    # 2.2 Grafico Linha
    ap(doc, "2.2 Grafico de Linha - Tendencias Anuais", sz=16, bold=True, after=12)
    ap(doc, (
        "Para visualizar as tendencias anuais, foi criado um grafico de "
        "linha com dois eixos: consumo total (kWh) no eixo esquerdo e valor "
        "medio da fatura (R$) no eixo direito."
    ), after=8)
    code(doc, """anual = df.groupby('ANO').agg(
    consumo_total=('CONSUMO_KWH','sum'),
    valor_medio=('VALOR','mean')
).reset_index()

fig, ax1 = plt.subplots(figsize=(10, 5))
ax1.plot(anual['ANO'], anual['consumo_total'],
         'o-', color='#2563eb', linewidth=2,
         label='Consumo Total (kWh)')
ax2 = ax1.twinx()
ax2.plot(anual['ANO'], anual['valor_medio'],
         's--', color='#dc2626', linewidth=2,
         label='Valor Medio Fatura (R$)')
plt.title('Tendencias Anuais')
plt.show()""")
    img(doc, 'grafico_linha_anual.png')
    doc.add_page_break()

    # 2.3 Barras
    ap(doc, "2.3 Grafico de Barras - Consumo Medio por Mes", sz=16, bold=True, after=12)
    ap(doc, (
        "O grafico de barras permite identificar padroes sazonais no consumo "
        "de energia. Barras em vermelho indicam meses com consumo acima da "
        "media, enquanto barras em azul ficam abaixo."
    ), after=8)
    code(doc, """mensal = df.groupby('MES')['CONSUMO_KWH'].mean()
mensal = mensal.reindex(meses_ordem)

colors = ['#ef4444' if v > mensal.mean()
          else '#3b82f6' for v in mensal]
plt.bar(mensal.index, mensal.values, color=colors)
plt.axhline(y=mensal.mean(), linestyle='--',
            color='gray', label='Media')
plt.title('Consumo Medio por Mes')
plt.show()""")
    img(doc, 'grafico_barras_mensal.png')
    doc.add_page_break()

    # 2.4 Correlacao
    ap(doc, "2.4 Analise de Correlacao", sz=16, bold=True, after=12)
    ap(doc, (
        "Foi realizada uma analise de correlacao entre o consumo (kWh) e o "
        "valor da fatura (R$) para verificar a relacao entre essas variaveis."
    ), after=8)
    code(doc, """corr = df['CONSUMO_KWH'].corr(df['VALOR'])
print(f'Correlacao: {corr:.2f}')

plt.scatter(df['CONSUMO_KWH'], df['VALOR'],
            c=df['ANO'], cmap='viridis', alpha=0.7)
plt.colorbar(label='Ano')
z = np.polyfit(df['CONSUMO_KWH'], df['VALOR'], 1)
plt.plot(x_sorted, np.poly1d(z)(x_sorted),
         '--', color='red', label='Tendencia')
plt.title(f'Correlacao: r = {corr:.2f}')
plt.show()""")
    img(doc, 'grafico_correlacao.png')

    ap(doc, (
        "O coeficiente de correlacao obtido foi de r = 0.62, indicando uma "
        "correlacao positiva moderada entre o consumo de energia e o valor "
        "da fatura. Isto significa que, em geral, quando o consumo aumenta, "
        "o valor da fatura tambem tende a aumentar, porem fatores como "
        "bandeiras tarifarias e reajustes influenciam essa relacao."
    ), after=12, before=8)
    doc.add_page_break()

    # 2.5 Plotly
    ap(doc, "2.5 Grafico Interativo com Plotly", sz=16, bold=True, after=12)
    ap(doc, (
        "Para enriquecer a analise com elementos interativos, foi criado um "
        "grafico utilizando a biblioteca Plotly, que permite zoom, hover com "
        "detalhes e selecao de series por ano."
    ), after=8)
    code(doc, """import plotly.express as px

fig = px.line(df, x=df.index, y='CONSUMO_KWH',
              color=df['ANO'].astype(str),
              title='Consumo Mensal - Visao Interativa')
fig.update_layout(template='plotly_white')
fig.show()""")
    img(doc, 'grafico_interativo.png')
    ap(doc, (
        "Os elementos interativos do Plotly permitem uma exploracao mais "
        "rica dos dados, facilitando a identificacao de outliers e padroes "
        "que nao seriam facilmente visiveis em graficos estaticos."
    ), after=12, before=8)
    doc.add_page_break()

    # ===== RESULTADOS =====
    ap(doc, "Resultados", sz=22, bold=True, after=16)
    ap(doc, "Respostas as perguntas propostas:", sz=14, bold=True, after=12)

    ap(doc, "1) Quais anos apresentaram as maiores reducoes no consumo?", sz=12, bold=True, after=6)
    ap(doc, (
        "Os anos de 2020 e 2021 apresentaram reducoes significativas no consumo "
        "energetico, possivelmente relacionadas a pandemia de COVID-19 e a "
        "reducao de atividades presenciais no Palacio do Jaburu."
    ), after=12)

    ap(doc, "2) Existe um padrao sazonal consistente?", sz=12, bold=True, after=6)
    ap(doc, (
        "Sim. Os meses mais quentes (setembro a dezembro) apresentam consumo "
        "consistentemente acima da media, enquanto meses como maio e junho "
        "tendem a ter consumo mais baixo, refletindo a menor necessidade de "
        "climatizacao no periodo de inverno."
    ), after=12)

    ap(doc, "3) Qual a relacao entre consumo e valor das faturas?", sz=12, bold=True, after=6)
    ap(doc, (
        "A correlacao positiva moderada (r = 0.62) indica que o aumento no "
        "consumo tende a elevar o valor da fatura, porem a relacao nao e "
        "perfeitamente linear devido a fatores como bandeiras tarifarias, "
        "reajustes anuais e impostos variados."
    ), after=12)

    ap(doc, "4) Como elementos interativos ajudam na analise?", sz=12, bold=True, after=6)
    ap(doc, (
        "Graficos interativos do Plotly permitem explorar os dados com zoom, "
        "hover e filtros por ano, facilitando a identificacao de outliers e "
        "padroes especificos que seriam dificeis de visualizar em graficos "
        "estaticos. A interatividade torna a narrativa de dados mais rica e "
        "acessivel para diferentes publicos."
    ), after=12)
    doc.add_page_break()

    # ===== CONCLUSAO =====
    ap(doc, "Conclusao", sz=22, bold=True, after=16)
    ap(doc, (
        "A pratica de visualizacao de dados e storytelling aplicada ao consumo "
        "energetico do Palacio do Jaburu demonstrou a importancia de utilizar "
        "graficos adequados para comunicar insights de forma clara e eficiente."
    ), after=12)
    ap(doc, (
        "As diferentes visualizacoes criadas — graficos de linha, barras, "
        "dispersao e interativos — permitiram identificar tendencias de "
        "reducao no consumo pos-2020, padroes sazonais com picos nos meses "
        "quentes, e uma correlacao positiva moderada entre consumo e custo."
    ), after=12)
    ap(doc, (
        "O uso combinado de Matplotlib para graficos estaticos e Plotly para "
        "visualizacoes interativas mostrou como ferramentas complementares "
        "enriquecem a narrativa baseada em dados, tornando a analise mais "
        "acessivel e a tomada de decisao mais fundamentada."
    ), after=12)

    # ===== REFERENCIAS =====
    ap(doc, "Referencias", sz=16, bold=True, after=12, before=24)
    ap(doc, "MCKINNEY, Wes. Python para Analise de Dados. 2. ed. Novatec, 2018.", after=6)
    ap(doc, "MATPLOTLIB. Matplotlib Documentation. Disponivel em: https://matplotlib.org/. Acesso em: maio 2026.", after=6)
    ap(doc, "PLOTLY. Plotly Python Documentation. Disponivel em: https://plotly.com/python/. Acesso em: maio 2026.", after=6)

    doc.save(DOCX)
    print(f"[OK] Documento gerado: {DOCX}")

if __name__ == "__main__":
    print("=" * 60)
    print("  GERADOR DE PORTFOLIO — Visualizacao de Dados")
    print("=" * 60)
    build()
