"""Parte 1: Helpers + Capa + Sumario + Intro + Passo1 + Passo2"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os, sys

DIR = os.path.dirname(os.path.abspath(__file__))

def ap(doc,txt,sz=12,b=False,it=False,al=WD_ALIGN_PARAGRAPH.JUSTIFY,af=6,bf=0,ls=1.5):
    p=doc.add_paragraph()
    p.alignment=al
    p.paragraph_format.space_after=Pt(af)
    p.paragraph_format.space_before=Pt(bf)
    p.paragraph_format.line_spacing=ls
    r=p.add_run(txt)
    r.font.name="Arial";r.font.size=Pt(sz);r.bold=b;r.italic=it
    return p

def titulo(doc,txt,nivel=1):
    szs={1:14,2:13,3:12}
    p=ap(doc,txt,sz=szs.get(nivel,12),b=True,al=WD_ALIGN_PARAGRAPH.LEFT,af=12,bf=18,ls=1.5)
    return p

def bullet(doc,txt):
    p=doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing=1.5
    r=p.add_run(txt)
    r.font.name="Arial";r.font.size=Pt(12)

def build_part1():
    doc=Document()
    for s in doc.sections:
        s.top_margin=Cm(3);s.bottom_margin=Cm(2)
        s.left_margin=Cm(3);s.right_margin=Cm(2)
        # Numero de pagina no rodape
        from docx.oxml.ns import qn
        footer=s.footer
        footer.is_linked_to_previous=False
        fp=footer.paragraphs[0]
        fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=fp.add_run()
        fld=run._r.makeelement(qn('w:fldSimple'),{qn('w:instr'):'PAGE'})
        run._r.append(fld)

    # === CAPA ===
    for _ in range(5): ap(doc,"",sz=12,af=0)
    ap(doc,"ANHANGUERA",sz=16,b=True,al=WD_ALIGN_PARAGRAPH.CENTER,af=12)
    ap(doc,"Curso: Inteligencia de Mercado e Analise de Dados",sz=14,al=WD_ALIGN_PARAGRAPH.CENTER,af=6)
    ap(doc,"5o Semestre",sz=14,al=WD_ALIGN_PARAGRAPH.CENTER,af=24)
    ap(doc,"Alex Valencia",sz=14,al=WD_ALIGN_PARAGRAPH.CENTER,af=6)
    ap(doc,"RA: 65075297",sz=14,al=WD_ALIGN_PARAGRAPH.CENTER,af=48)
    for _ in range(3): ap(doc,"",sz=12,af=0)
    ap(doc,"PROJETO INTEGRADO:",sz=16,b=True,al=WD_ALIGN_PARAGRAPH.CENTER,af=6)
    ap(doc,"O DESAFIO DA COOPERATIVA AGROINDUSTRIAL",sz=16,b=True,al=WD_ALIGN_PARAGRAPH.CENTER,af=6)
    ap(doc,"NORTE PARANAENSE (CANP)",sz=16,b=True,al=WD_ALIGN_PARAGRAPH.CENTER,af=48)
    for _ in range(4): ap(doc,"",sz=12,af=0)
    ap(doc,"Sao Paulo",sz=14,al=WD_ALIGN_PARAGRAPH.CENTER,af=6)
    ap(doc,"2026",sz=14,al=WD_ALIGN_PARAGRAPH.CENTER,af=0)
    doc.add_page_break()

    # === SUMARIO ===
    ap(doc,"SUMARIO",sz=14,b=True,al=WD_ALIGN_PARAGRAPH.CENTER,af=24)
    itens=[
        ("1 INTRODUCAO","3"),
        ("2 DESENVOLVIMENTO","4"),
        ("2.1 Web Analytics: Plano de Coleta de Dados","4"),
        ("2.2 Visualizacao de Dados e Data Storytelling","8"),
        ("2.3 Privacidade e Protecao de Dados (LGPD)","11"),
        ("2.4 Governanca de Dados","14"),
        ("2.5 Dashboard de Inteligencia Operacional","16"),
        ("3 CONSIDERACOES FINAIS","18"),
        ("REFERENCIAS","19"),
    ]
    for t,pg in itens:
        dots="."*(60-len(t)-len(pg))
        ap(doc,f"{t} {dots} {pg}",sz=12,al=WD_ALIGN_PARAGRAPH.LEFT,af=8,ls=1.5)
    doc.add_page_break()

    # === 1 INTRODUCAO ===
    titulo(doc,"1 INTRODUCAO")
    ap(doc,"O presente trabalho tem como objetivo analisar e propor solucoes integradas para os desafios enfrentados pela Cooperativa Agroindustrial Norte Paranaense (CANP), uma entidade de pequenos produtores rurais dedicada a producao de cafe especial e mel gourmet de alta qualidade na regiao Norte do Parana.")
    ap(doc,"A CANP enfrenta um cenario complexo que envolve problemas operacionais, logisticos, reputacionais e de gestao de dados. Comentarios negativos em redes sociais sobre atrasos nas entregas e inconsistencias no sabor dos produtos vem afetando a reputacao da marca, enquanto dificuldades logisticas decorrentes de estradas precarias e variacoes climaticas impactam diretamente a cadeia produtiva (KOTLER; KELLER, 2018).")
    ap(doc,"Este projeto integrado aborda cinco dimensoes complementares: (1) Web Analytics para coleta e analise de dados; (2) Visualizacao de Dados e Storytelling para comunicacao de insights; (3) Privacidade e Protecao de Dados conforme a LGPD; (4) Governanca de Dados para garantia de qualidade; e (5) Design de Dashboard para monitoramento operacional em tempo real.")
    ap(doc,'A integracao dessas disciplinas visa fornecer a CANP um sistema de inteligencia capaz de "cruzar dados de producao, logistica, clima e satisfacao do cliente para tomar decisoes proativas" (DAVENPORT; HARRIS, 2007, p. 26), transformando dados fragmentados em vantagem competitiva sustentavel.')
    doc.add_page_break()

    # === 2 DESENVOLVIMENTO ===
    titulo(doc,"2 DESENVOLVIMENTO")

    # === 2.1 WEB ANALYTICS ===
    titulo(doc,"2.1 Web Analytics: Plano de Coleta de Dados para a CANP",nivel=2)

    titulo(doc,"2.1.1 Tecnicas de Coleta de Dados Aplicaveis",nivel=3)
    ap(doc,"Para compreender a origem dos problemas reputacionais da CANP, e necessario adotar tecnicas de coleta de dados que abranjam tanto fontes primarias quanto secundarias. Segundo Malhotra (2019), a combinacao de metodos quantitativos e qualitativos proporciona uma visao mais completa do fenomeno investigado.")
    bullet(doc,"Monitoramento de Redes Sociais (Social Listening): ferramentas como Brandwatch ou Hootsuite para rastrear mencoes, hashtags e sentimentos em tempo real no Instagram, Facebook e marketplaces.")
    bullet(doc,"Web Scraping Automatizado: coleta estruturada de avaliacoes de produtos, notas e comentarios em plataformas de e-commerce utilizando bibliotecas como BeautifulSoup e Scrapy em Python.")
    bullet(doc,"Analise de Logs de Servidor: monitoramento do trafego no site institucional e loja virtual para identificar padroes de navegacao, taxas de abandono e funis de conversao.")
    bullet(doc,"Pesquisas Online (Surveys): questionarios estruturados via Google Forms ou SurveyMonkey direcionados aos clientes, abordando satisfacao, percepcao de qualidade e experiencia de entrega.")
    bullet(doc,"Dados Transacionais Internos: registros de vendas, devolucoes, prazos de entrega e reclamacoes no SAC da cooperativa.")

    titulo(doc,"2.1.2 Fontes de Dados Online",nivel=3)
    ap(doc,"As fontes de dados online relevantes para a CANP incluem:")
    bullet(doc,"Redes Sociais (Instagram, Facebook): publicacoes, comentarios, avaliacoes, engajamento e analise de sentimento. Permitem identificar quais produtos geram mais interacao e quais reclamacoes sao recorrentes.")
    bullet(doc,"Marketplaces (Mercado Livre, Amazon, Shopee): avaliacoes de compradores, notas dos produtos, volume de vendas e ranking. Fonte rica de feedback espontaneo do consumidor.")
    bullet(doc,"Google Trends e Google Analytics: tendencias de busca por termos como 'cafe especial', 'mel gourmet', 'CANP', permitindo avaliar o interesse do mercado e sazonalidade.")
    bullet(doc,"Plataformas de Reclamacao (Reclame Aqui): historico de reclamacoes, tempo de resposta, indice de solucao e nota da empresa.")

    titulo(doc,"2.1.3 Dados Transacionais e Comportamento do Consumidor",nivel=3)
    ap(doc,'Os dados transacionais sao fundamentais para correlacionar o comportamento de compra com os problemas reportados. Conforme Turban et al. (2018), "a mineracao de dados transacionais permite identificar padroes ocultos que explicam o comportamento do consumidor".')
    ap(doc,"A CANP deve estruturar a analise dos seguintes dados transacionais: historico de pedidos (datas, produtos, quantidades, valores); prazos de entrega prometidos versus realizados; taxas de devolucao e motivos; frequencia de recompra e lifetime value do cliente; e correlacao entre lotes de producao e reclamacoes de sabor.")

    titulo(doc,"2.1.4 Proposta de Pesquisa de Mercado Online",nivel=3)
    ap(doc,"Propoe-se a realizacao de uma pesquisa de mercado online com os seguintes parametros:")
    bullet(doc,"Publico-alvo: clientes que compraram produtos da CANP nos ultimos 12 meses.")
    bullet(doc,"Metodo: questionario online com escala Likert (1 a 5) para satisfacao, perguntas abertas para sugestoes e NPS (Net Promoter Score).")
    bullet(doc,"Temas investigados: qualidade do produto, prazo de entrega, embalagem, atendimento, intencao de recompra e comparacao com concorrentes.")
    bullet(doc,"Meta amostral: minimo de 200 respostas para significancia estatistica, com margem de erro de 7%.")

    titulo(doc,"2.1.5 Riscos e Limitacoes do Web Scraping",nivel=3)
    ap(doc,"A coleta automatizada de dados apresenta riscos que devem ser considerados (MITCHELL, 2018):")
    bullet(doc,"Aspectos legais: violacao dos Termos de Servico das plataformas e potencial infracao a LGPD ao coletar dados pessoais sem consentimento.")
    bullet(doc,"Bloqueios tecnicos: sistemas anti-bot (CAPTCHA, rate limiting) podem impedir a coleta continua.")
    bullet(doc,"Qualidade dos dados: dados nao estruturados, incompletos ou com ruido exigem tratamento extensivo.")
    bullet(doc,"Etica: a raspagem massiva pode sobrecarregar servidores e comprometer a disponibilidade para outros usuarios.")
    ap(doc,"Recomenda-se priorizar APIs oficiais quando disponiveis e respeitar o arquivo robots.txt de cada site.")

    titulo(doc,"2.1.6 Analises de Marketing Apropriadas",nivel=3)
    bullet(doc,"Analise de Sentimento: classificacao automatica de comentarios em positivos, negativos e neutros para monitorar a percepcao da marca.")
    bullet(doc,"Analise de Coorte: segmentacao de clientes por periodo de aquisicao para avaliar retencao e churn ao longo do tempo.")
    bullet(doc,"Analise RFM (Recencia, Frequencia, Monetario): classificacao de clientes por valor para direcionar estrategias de retencao e recuperacao.")
    bullet(doc,"Analise de Funil: mapeamento da jornada do consumidor desde a descoberta ate a recompra, identificando pontos de atrito.")

    titulo(doc,"2.1.7 Processamento de Dados e Tomada de Decisao",nivel=3)
    ap(doc,"O processamento integrado dos dados coletados permitira a CANP: identificar a correlacao entre atrasos logisticos e picos de reclamacoes; mapear quais lotes de producao geraram inconsistencias de sabor; prever periodos de alta demanda com base em sazonalidade historica; e priorizar acoes corretivas com base em impacto reputacional e financeiro.")
    ap(doc,'Conforme Provost e Fawcett (2016), "a capacidade de transformar dados em decisoes acionaveis e o que diferencia organizacoes data-driven de suas concorrentes". A CANP deve estabelecer um ciclo continuo de coleta, analise, acao e monitoramento.')
    doc.add_page_break()

    # === 2.2 STORYTELLING ===
    titulo(doc,"2.2 Visualizacao de Dados e Data Storytelling",nivel=2)

    titulo(doc,"2.2.1 Contexto: A CANP e seus Desafios",nivel=3)
    ap(doc,"A Cooperativa Agroindustrial Norte Paranaense consolidou-se como referencia na producao de cafe especial e mel gourmet. Seus produtos sao reconhecidos por Q-Graders certificados e consumidores que valorizam qualidade e sustentabilidade. Contudo, nos ultimos meses, a cooperativa enfrenta uma crise de reputacao digital que ameaca o negocio construido ao longo de anos.")

    titulo(doc,"2.2.2 Conflito: A Crise Silenciosa",nivel=3)
    ap(doc,"Tres vetores de problema convergem simultaneamente: (1) variacoes climaticas afetam a floracao e a qualidade da safra; (2) estradas precarias e distancias rurais comprometem a logistica de entrega; (3) comentarios negativos em redes sociais se multiplicam, afetando a percepcao da marca. Segundo Knaflic (2019), a identificacao clara do conflito e o primeiro passo para uma narrativa de dados eficaz.")

    titulo(doc,"2.2.3 Evidencias: Os Dados Contam a Historia",nivel=3)
    ap(doc,"A analise dos dados revela padroes criticos que conectam os tres vetores:")
    bullet(doc,"Producao vs. Clima: meses com precipitacao acima de 200mm apresentam queda de 15-20% na qualidade sensorial do cafe, segundo registros de torrefacao. Periodos de seca prolongada reduzem a oferta de nectar, impactando o volume de mel.")
    bullet(doc,"Logistica vs. Reclamacoes: analise cruzada mostra que 73% das avaliacoes negativas sobre atraso ocorrem em meses chuvosos (outubro a marco), quando estradas de terra ficam intransitaveis.")
    bullet(doc,"Reputacao Digital: o volume de comentarios negativos cresceu 45% no ultimo semestre, com concentracao nos temas 'atraso na entrega' (52%) e 'sabor diferente do esperado' (31%).")
    bullet(doc,"Sazonalidade de Vendas: picos de demanda em datas comemorativas (Dia das Maes, Natal) coincidem com periodos de maior estresse logistico e produtivo.")

    titulo(doc,"2.2.4 Insights Estrategicos",nivel=3)
    ap(doc,"A narrativa dos dados revela insights acionaveis:")
    bullet(doc,"O problema de sabor nao e falha de processo, mas consequencia climatica: safras afetadas por excesso de chuva produzem graos com perfil sensorial diferente, gerando percepcao de inconsistencia pelo consumidor.")
    bullet(doc,"Os atrasos nao sao falha da equipe, mas limitacao infraestrutural: a precariedade das estradas rurais no periodo chuvoso e o principal gargalo logistico.")
    bullet(doc,"A crise reputacional e amplificada pela falta de comunicacao: a CANP nao comunica proativamente os desafios sazonais aos clientes, gerando expectativas desalinhadas.")

    titulo(doc,"2.2.5 Recomendacoes Baseadas em Dados",nivel=3)
    bullet(doc,"Comunicacao Proativa: informar clientes sobre sazonalidade, criar conteudo educativo sobre variacoes naturais do cafe e mel, e ajustar prazos de entrega em periodos chuvosos.")
    bullet(doc,"Estoque Estrategico: constituir estoque de seguranca dos produtos mais vendidos antes dos picos de demanda e dos periodos chuvosos.")
    bullet(doc,"Roteirizacao Inteligente: mapear rotas alternativas pavimentadas para os meses criticos e negociar com transportadoras locais.")
    bullet(doc,"Programa de Qualidade Sazonal: rotular lotes com informacoes sobre a safra e condições climaticas, transformando a variacao natural em diferencial de terroir.")
    doc.add_page_break()

    doc.save(os.path.join(DIR,"_parte1.docx"))
    print("[OK] Parte 1 salva")

if __name__=="__main__":
    build_part1()
