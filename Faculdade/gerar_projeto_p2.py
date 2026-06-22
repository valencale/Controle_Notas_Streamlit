"""Parte 2: Passo3 (LGPD) + Passo4 (Governanca) + Passo5 (Dashboard) + Conclusao + Refs"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DIR = os.path.dirname(os.path.abspath(__file__))

def ap(doc,txt,sz=12,b=False,it=False,al=WD_ALIGN_PARAGRAPH.JUSTIFY,af=6,bf=0,ls=1.5):
    p=doc.add_paragraph();p.alignment=al
    p.paragraph_format.space_after=Pt(af);p.paragraph_format.space_before=Pt(bf);p.paragraph_format.line_spacing=ls
    r=p.add_run(txt);r.font.name="Arial";r.font.size=Pt(sz);r.bold=b;r.italic=it
    return p

def titulo(doc,txt,nivel=1):
    szs={1:14,2:13,3:12}
    return ap(doc,txt,sz=szs.get(nivel,12),b=True,al=WD_ALIGN_PARAGRAPH.LEFT,af=12,bf=18,ls=1.5)

def bullet(doc,txt):
    p=doc.add_paragraph(style="List Bullet");p.paragraph_format.line_spacing=1.5
    r=p.add_run(txt);r.font.name="Arial";r.font.size=Pt(12)

def build_part2():
    # Abrir parte 1
    doc=Document(os.path.join(DIR,"_parte1.docx"))

    # === 2.3 LGPD ===
    titulo(doc,"2.3 Privacidade e Protecao de Dados",nivel=2)

    titulo(doc,"2.3.1 Inventario de Dados (Data Mapping) para a CANP",nivel=3)
    ap(doc,"Conforme o artigo 37 da LGPD (Lei 13.709/2018), todo controlador deve manter registro das operacoes de tratamento de dados pessoais. Para a CANP, o inventario deve ser estruturado considerando as particularidades do ambiente rural e a diversidade de fontes de dados (BIONI, 2019).")

    ap(doc,"Estrutura do Inventario de Dados:",b=True,af=6,bf=12)
    bullet(doc,"Dados de Cooperados: nome, CPF, endereco da propriedade, dados bancarios, producao por safra. Base legal: execucao de contrato (Art. 7, V).")
    bullet(doc,"Dados de Producao: registros de colheita, temperatura de torra, umidade, tipo floral do mel, rastreabilidade de lotes. Classificacao: dados nao pessoais (operacionais).")
    bullet(doc,"Dados Climaticos: precipitacao, temperatura, umidade do ar coletados por sensores IoT nas propriedades. Classificacao: dados nao pessoais.")
    bullet(doc,"Dados de Clientes: nome, e-mail, endereco de entrega, historico de compras, preferencias. Base legal: consentimento (Art. 7, I) ou execucao de contrato.")
    bullet(doc,"Dados de Redes Sociais: comentarios publicos, avaliacoes, perfis de usuarios. Base legal: interesse legitimo (Art. 7, IX), desde que respeitada a expectativa do titular.")

    ap(doc,"Principais Riscos de Privacidade:",b=True,af=6,bf=12)
    bullet(doc,"Coleta: aplicativos nao padronizados utilizados pelos cooperados podem coletar dados excessivos ou sem consentimento adequado, violando o principio da necessidade.")
    bullet(doc,"Transmissao: a conectividade limitada nas areas rurais forca o uso de meios inseguros (WhatsApp, e-mail pessoal), sem criptografia adequada.")
    bullet(doc,"Armazenamento: planilhas locais em computadores sem backup ou controle de acesso expoe os dados a perdas e acessos nao autorizados.")
    bullet(doc,"Uso: ausencia de politica clara sobre quem pode acessar cada tipo de dado, permitindo uso indevido ou compartilhamento nao autorizado.")
    bullet(doc,"Descarte: sem procedimento de descarte seguro, dados antigos de clientes e cooperados permanecem armazenados indefinidamente.")

    ap(doc,"Controles Recomendados:",b=True,af=6,bf=12)
    bullet(doc,"Tecnicos: implantacao de VPN para transmissao segura; criptografia de banco de dados; backup automatizado em nuvem; autenticacao multifator para sistemas criticos; aplicativo padronizado para coleta em campo.")
    bullet(doc,"Administrativos: nomeacao de Encarregado de Dados (DPO); treinamento anual dos cooperados sobre LGPD; politica de privacidade publicada no site; termos de consentimento para clientes; procedimento de resposta a incidentes.")
    doc.add_page_break()

    titulo(doc,"2.3.2 Tratamento de Dados de Consumidores e Reputacao Digital",nivel=3)
    ap(doc,"A CANP opera em ambientes digitais que coletam dados pessoais dos consumidores de forma continua. A analise dos riscos e a aplicacao dos principios da LGPD sao essenciais para proteger tanto os titulares quanto a propria cooperativa (PINHEIRO, 2020).")

    ap(doc,"Riscos Identificados:",b=True,af=6,bf=12)
    bullet(doc,"Coleta excessiva: formularios de cadastro que solicitam dados alem do necessario para a transacao (ex: data de nascimento, genero) sem justificativa clara.")
    bullet(doc,"Compartilhamento com terceiros: integracao com marketplaces e redes sociais pode resultar em compartilhamento de dados sem ciencia do titular.")
    bullet(doc,"Perfilamento sem consentimento: uso de cookies e rastreamento de navegacao para segmentacao publicitaria sem aviso previo.")
    bullet(doc,"Exposicao em respostas publicas: ao responder reclamacoes em redes sociais, a CANP pode expor dados do cliente (numero de pedido, endereco) publicamente.")

    ap(doc,"Aplicacao dos Principios da LGPD:",b=True,af=6,bf=12)
    bullet(doc,"Finalidade: coletar apenas dados necessarios para a finalizacao da compra e comunicacao pos-venda. Qualquer uso adicional (marketing, pesquisa) requer consentimento especifico.")
    bullet(doc,"Necessidade: revisar formularios para eliminar campos desnecessarios, adotando o principio da minimizacao de dados.")
    bullet(doc,"Adequacao: garantir que o tratamento dos dados seja compativel com a finalidade informada ao titular no momento da coleta.")
    bullet(doc,"Prevencao: implementar medidas de seguranca proativas, como criptografia, controle de acesso e monitoramento de vazamentos.")
    bullet(doc,"Responsabilizacao: documentar todas as operacoes de tratamento, manter registros de consentimento e estar preparado para demonstrar conformidade a ANPD.")

    ap(doc,"Boas Praticas para Marketing, SAC e Gestao de Reputacao:",b=True,af=6,bf=12)
    bullet(doc,"Marketing: opt-in explicito para newsletters; opcao de descadastro em todas as comunicacoes; segmentacao baseada em dados anonimizados.")
    bullet(doc,"SAC: tratar reclamacoes em canais privados (DM, e-mail) para evitar exposicao de dados; treinar equipe sobre o que pode ser compartilhado publicamente.")
    bullet(doc,"Reputacao: monitorar avaliacoes sem coletar dados pessoais adicionais; responder publicamente de forma generica e direcionar para canais privados.")
    doc.add_page_break()

    # === 2.4 GOVERNANCA ===
    titulo(doc,"2.4 Governanca de Dados: Plano de Qualidade de Dados",nivel=2)
    ap(doc,'A governanca de dados e definida por Ladley (2019) como "o exercicio de autoridade, controle e tomada de decisao compartilhada sobre a gestao de ativos de dados". Para a CANP, a implementacao de um Plano de Qualidade de Dados (PQD) e urgente.')

    titulo(doc,"2.4.1 Dimensoes de Qualidade para o Dataset de Colheita de Cafe",nivel=3)
    bullet(doc,"Completude: todos os campos obrigatorios devem estar preenchidos (data de colheita, propriedade, volume, variedade, metodo de processamento). Meta: 95% de registros completos.")
    bullet(doc,"Consistencia: padronizacao de unidades (kg, sacas de 60kg), formatos de data (DD/MM/AAAA) e nomenclatura de variedades. Registros conflitantes entre fontes devem ser sinalizados.")
    bullet(doc,"Acuracia: valores devem refletir a realidade. Peso registrado manualmente deve ser validado contra pesagem mecanica. Desvios acima de 5% devem gerar alerta.")
    bullet(doc,"Temporalidade: dados devem ser registrados em ate 24 horas apos a colheita. Registros com atraso superior a 72 horas devem ser marcados como 'nao confiavel'.")
    bullet(doc,"Unicidade: cada lote de colheita deve ter um identificador unico. Registros duplicados devem ser automaticamente detectados e eliminados.")
    bullet(doc,"Validade: valores devem respeitar faixas aceitas (ex: umidade entre 10-14% para cafe beneficiado; peso por saca entre 55-65kg).")

    titulo(doc,"2.4.2 Checklist de Limpeza de Dados",nivel=3)
    bullet(doc,"Identificar e tratar valores nulos: substituir por mediana (numericos) ou 'NAO INFORMADO' (categoricos) com flag indicativo.")
    bullet(doc,"Padronizar unidades de medida: converter todas as medidas para kg e litros, documentando fatores de conversao.")
    bullet(doc,"Normalizar nomes de propriedades e variedades: criar tabela de referencia (dicionario de dados) e aplicar mapeamento.")
    bullet(doc,"Detectar outliers: aplicar metodo IQR (Interquartile Range) para identificar valores atipicos em peso e umidade.")
    bullet(doc,"Remover duplicatas: implementar verificacao por chave composta (data + propriedade + lote).")
    bullet(doc,"Validar integridade referencial: garantir que cada registro de colheita referencia uma propriedade e cooperado validos.")
    bullet(doc,"Documentar todas as transformacoes: manter log de alteracoes com data, responsavel e justificativa.")

    titulo(doc,"2.4.3 Data Owners por Dataset Critico",nivel=3)
    bullet(doc,"Dataset Clima: Engenheiro Agronomo da CANP. Responsavel por validar dados de sensores IoT, garantir calibracao dos equipamentos e definir alertas de condicoes adversas.")
    bullet(doc,"Dataset Producao (Colheita e Processamento): Gerente de Producao. Responsavel por garantir registro tempestivo, padronizacao de variedades e rastreabilidade dos lotes.")
    bullet(doc,"Dataset Logistica: Coordenador de Logistica. Responsavel por dados de frota, rotas, tempos de entrega, custos de transporte e condicoes das estradas.")
    ap(doc,'Cada data owner deve ter autonomia para definir regras de qualidade, aprovar acessos e responder por inconsistencias em seu dominio, conforme recomendado pelo DAMA-DMBOK (DAMA International, 2017).')
    doc.add_page_break()

    # === 2.5 DASHBOARD ===
    titulo(doc,"2.5 Dashboard de Inteligencia Operacional e de Mercado (DIOM)",nivel=2)

    titulo(doc,"2.5.1 Fontes de Dados Cruciais",nivel=3)
    ap(doc,"Fonte 1 - Sistema de Producao e Clima:",b=True,af=4,bf=8)
    ap(doc,"Integra dados de sensores IoT (temperatura, umidade, precipitacao) com registros de colheita, processamento e torrefacao. Justificativa: permite correlacionar condicoes climaticas com qualidade da safra e antecipar impactos na producao.")
    ap(doc,"Fonte 2 - Plataforma de Logistica e Vendas:",b=True,af=4,bf=8)
    ap(doc,"Consolida dados de pedidos, prazos de entrega, status de transporte e devolucoes. Justificativa: possibilita identificar gargalos logisticos em tempo real e avaliar o cumprimento dos SLAs de entrega.")
    ap(doc,"Fonte 3 - Monitoramento de Reputacao Digital:",b=True,af=4,bf=8)
    ap(doc,"Agrega avaliacoes de redes sociais, marketplaces e Reclame Aqui via APIs e social listening. Justificativa: transforma o sentimento do consumidor em indicador quantificavel para decisao rapida.")

    titulo(doc,"2.5.2 KPIs por Area Critica",nivel=3)
    ap(doc,"KPI de Producao - Indice de Qualidade Sensorial (IQS):",b=True,af=4,bf=8)
    ap(doc,"Definicao: percentual de lotes aprovados na avaliacao sensorial (nota >= 80 SCA). Meta: >= 85%. Acao: se o IQS cair abaixo de 80%, acionar auditoria no processo de torra e verificar correlacao com dados climaticos da safra. Revisar calibracao dos equipamentos de torrefacao.")
    ap(doc,"KPI de Logistica - Taxa de Entrega no Prazo (TEP):",b=True,af=4,bf=8)
    ap(doc,"Definicao: percentual de pedidos entregues dentro do prazo prometido. Meta: >= 92%. Acao: se a TEP cair abaixo de 85%, ativar rotas alternativas pavimentadas, acionar transportadoras reservas e comunicar proativamente os clientes sobre possiveis atrasos.")
    ap(doc,"KPI de Mercado/Reputacao - Indice de Sentimento Digital (ISD):",b=True,af=4,bf=8)
    ap(doc,"Definicao: razao entre avaliacoes positivas e o total de avaliacoes nas plataformas digitais. Meta: >= 0.75 (75% positivas). Acao: se o ISD cair abaixo de 0.65, acionar equipe de SAC para resposta prioritaria, investigar causa-raiz nas areas de producao e logistica, e publicar comunicado institucional.")

    titulo(doc,"2.5.3 Visualizacao e Tomada de Decisao",nivel=3)
    ap(doc,"O dashboard deve apresentar os tres KPIs em painel unico com semaforo visual (verde/amarelo/vermelho), graficos de tendencia dos ultimos 30 dias e alertas automaticos por e-mail quando qualquer indicador ultrapassar o limiar critico. A visualizacao integrada permite ao gestor identificar rapidamente a origem dos problemas e acionar a area responsavel, conforme o principio de management by exception (FEW, 2013).")
    doc.add_page_break()

    # === 3 CONSIDERACOES FINAIS ===
    titulo(doc,"3 CONSIDERACOES FINAIS")
    ap(doc,"O presente projeto integrado demonstrou que os desafios enfrentados pela CANP nao sao isolados, mas interconectados em uma cadeia que vai do clima a reputacao digital. A aplicacao integrada das cinco disciplinas estudadas revelou que a cooperativa possui dados suficientes para tomar decisoes proativas, porem carece de estrutura para coleta-los, organiza-los e transforma-los em insights acionaveis.")
    ap(doc,"O plano de Web Analytics proposto fornece a base para coleta estruturada de dados internos e externos. O Data Storytelling traduz esses dados em narrativas convincentes para a diretoria. A conformidade com a LGPD protege a cooperativa juridicamente e fortalece a confianca dos consumidores. A governanca de dados garante qualidade e confiabilidade das informacoes. E o dashboard operacional integra tudo em uma ferramenta de monitoramento e acao em tempo real.")
    ap(doc,"A implementacao dessas solucoes exige investimento em tecnologia, capacitacao e mudanca cultural, mas o retorno esperado em eficiencia operacional, satisfacao do cliente e protecao reputacional justifica amplamente o esforco. A CANP tem a oportunidade de se tornar referencia nao apenas em qualidade de produtos, mas tambem em gestao inteligente de dados no setor agroindustrial.")
    doc.add_page_break()

    # === REFERENCIAS ===
    titulo(doc,"REFERENCIAS")
    refs = [
        "BIONI, Bruno Ricardo. Protecao de Dados Pessoais: a funcao e os limites do consentimento. Rio de Janeiro: Forense, 2019.",
        "BRASIL. Lei n. 13.709, de 14 de agosto de 2018. Lei Geral de Protecao de Dados Pessoais (LGPD). Diario Oficial da Uniao, Brasilia, DF, 15 ago. 2018.",
        "DAMA INTERNATIONAL. DAMA-DMBOK: Data Management Body of Knowledge. 2. ed. Basking Ridge: Technics Publications, 2017.",
        "DAVENPORT, Thomas H.; HARRIS, Jeanne G. Competing on Analytics: The New Science of Winning. Boston: Harvard Business School Press, 2007.",
        "FEW, Stephen. Information Dashboard Design: Displaying Data for At-a-Glance Monitoring. 2. ed. Burlingame: Analytics Press, 2013.",
        "KNAFLIC, Cole Nussbaumer. Storytelling com Dados: um guia sobre visualizacao de dados para profissionais de negocios. Rio de Janeiro: Alta Books, 2019.",
        "KOTLER, Philip; KELLER, Kevin Lane. Administracao de Marketing. 15. ed. Sao Paulo: Pearson, 2018.",
        "LADLEY, John. Data Governance: How to Design, Deploy, and Sustain an Effective Data Governance Program. 2. ed. London: Academic Press, 2019.",
        "MALHOTRA, Naresh K. Pesquisa de Marketing: uma orientacao aplicada. 7. ed. Porto Alegre: Bookman, 2019.",
        "MITCHELL, Ryan. Web Scraping with Python. 2. ed. Sebastopol: O'Reilly Media, 2018.",
        "PINHEIRO, Patricia Peck. Protecao de Dados Pessoais: comentarios a Lei n. 13.709/2018 (LGPD). 2. ed. Sao Paulo: Saraiva, 2020.",
        "PROVOST, Foster; FAWCETT, Tom. Data Science para Negocios. Rio de Janeiro: Alta Books, 2016.",
        "TURBAN, Efraim et al. Business Intelligence, Analytics, and Data Science: A Managerial Perspective. 4. ed. New York: Pearson, 2018.",
    ]
    for ref in refs:
        ap(doc,ref,sz=12,al=WD_ALIGN_PARAGRAPH.LEFT,af=12,ls=1.0)

    # SALVAR FINAL
    final=os.path.join(DIR,"Projeto_Integrado_CANP_Alex_Valencia.docx")
    doc.save(final)
    print(f"[OK] Documento final: {final}")
    # Limpar temp
    try: os.remove(os.path.join(DIR,"_parte1.docx"))
    except: pass

if __name__=="__main__":
    print("="*60)
    print("  PROJETO INTEGRADO — CANP")
    print("="*60)
    build_part2()
