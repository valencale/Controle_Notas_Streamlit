import os
import base64
import json
import urllib.request
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_mermaid_image(doc, mermaid_code, filename):
    try:
        json_data = json.dumps({"code": mermaid_code, "mermaid": {"theme": "default"}})
        base64_data = base64.b64encode(json_data.encode('utf-8')).decode('utf-8')
        url = f"https://mermaid.ink/img/{base64_data}"
        
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response, open(filename, 'wb') as out_file:
            data = response.read()
            out_file.write(data)
        
        doc.add_picture(filename, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        os.remove(filename)
    except Exception as e:
        print(f"Erro ao gerar imagem mermaid: {e}")
        doc.add_paragraph("[Diagrama de Fluxo - Não foi possível gerar a imagem online]")
        doc.add_paragraph(mermaid_code)

def create_doc():
    doc = Document()
    
    # Capa
    title = doc.add_heading('Controle de Notas\nManual Completo e Guia de Arquitetura', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    
    # ---------------------------------------------------------
    # PARTE 1: GUIA DE INÍCIO RÁPIDO E USO (Foco no Usuário)
    # ---------------------------------------------------------
    add_heading(doc, 'PARTE 1: Guia de Início Rápido e Uso', 1, RGBColor(33, 150, 243))
    doc.add_paragraph('Bem-vindo ao Controle de Notas! Este módulo da documentação é focado no uso diário do sistema. Usaremos uma linguagem progressiva, guiando-o pelas principais etapas lógicas que movem um pedido do faturamento até a entrega.')

    # Introdução
    add_heading(doc, '1.1 O que é o Sistema?', 2)
    doc.add_paragraph('O Controle de Notas consolida dados de faturamento em um ambiente rápido e visual. Ele diminui o esforço de digitação, minimiza erros de conferência e oferece inteligência de roteirização geográfica para entregas.')

    # Início Rápido
    add_heading(doc, '1.2 Guia Rápido (Quick Start)', 2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Ligar o Sistema: ').bold = True
    p.add_run('Vá até a pasta do projeto e dê um duplo clique em ')
    p.add_run('START_APP.bat').italic = True
    p.add_run('. O seu navegador abrirá a tela principal automaticamente.')
    
    # Uso Progressivo
    add_heading(doc, '1.3 Manual Passo a Passo', 2)

    doc.add_heading('Passo 1: Ingestão de Dados (PDF)', 3)
    doc.add_paragraph('Acesse o menu lateral e vá em "Ingestão PDF". O sistema lê o formato do "Mapa de Separação por Pedido" (Onfinity). Você só precisa arrastar os arquivos PDF para a tela. Em segundos, o sistema extrai os números dos pedidos, endereços, observações e insere na Esteira de Logística automaticamente.')

    doc.add_heading('Passo 2: Gestão de Esteira', 3)
    doc.add_paragraph('Na página "Esteira", visualize a tabela principal onde cada cor representa um "Status" (Em Separação, Faturado, Em Rota). Use as barras de busca para filtrar por número ou nome do cliente. Se algo der errado na separação, clique no status do item para modificá-lo.')

    doc.add_heading('Passo 3: Checklist Multi-picking', 3)
    doc.add_paragraph('Na aba "Checklist Separação", selecione os pedidos de uma determinada carga e um checklist único é gerado. Vá marcando as caixinhas (checkboxes) conforme coleta os itens no estoque. Isso garante 100% de precisão antes de o caminhão ser carregado.')

    doc.add_heading('Passo 4: Mapa e Expedição', 3)
    doc.add_paragraph('Com a carga separada, use a aba "Mapa" para analisar os locais de entrega em um mapa interativo. Depois, na aba "Expedição", crie um novo romaneio informando a data e o motorista. Os itens sairão da "Esteira" e migrarão para o status "Em Rota".')

    doc.add_heading('Passo 5: Histórico e Analytics', 3)
    doc.add_paragraph('Finalizou o dia? Visite a tela de "Histórico" ou "Histórico de Expedições" para auditar pedidos já entregues. Se a diretoria precisar de resultados, o "BI Analytics" exibe gráficos instantâneos da performance geral de expedição.')

    doc.add_page_break()

    # ---------------------------------------------------------
    # PARTE 2: ARQUITETURA E API (Foco Técnico)
    # ---------------------------------------------------------
    add_heading(doc, 'PARTE 2: Arquitetura e Estrutura Técnica', 1, RGBColor(33, 150, 243))
    doc.add_paragraph('Esta seção apresenta os fundamentos lógicos (design patterns), dependências do ecossistema Python e como a aplicação conversa internamente.')

    # Diagramas 
    add_heading(doc, '2.1 Fluxogramas e Ciclo de Vida', 2)
    
    doc.add_heading('A. Ciclo de Dados Principal (Backend Excel)', 3)
    arch_code = """sequenceDiagram
    participant UI as Streamlit App
    participant BE as excel_handler.py
    participant DB as CONTROLE NOTAS.xlsm
    
    UI->>BE: insert_pedido() / update_status()
    BE->>DB: Open workbook (keep_vba=True)
    BE->>DB: Modify Cell/Row
    BE->>DB: Save & Close
    BE->>UI: cache_data.clear()
    UI->>UI: Atualiza Tabela de Estado"""
    add_mermaid_image(doc, arch_code, 'arch_cycle.png')
    doc.add_paragraph('Padrão: Open-Modify-Save-Close garante mínimo de concorrência com acessos humanos simultâneos à planilha. O Streamlit Cache previne leituras repetidas desnecessárias.')

    doc.add_heading('B. Ciclo de Extração de PDF', 3)
    pdf_arch_code = """graph TD
    A[Arquivo PDF (Onfinity)] --> B[pdf_parser.py]
    B --> C{É Página de Continuação?}
    C -->|Sim| D[Mescla Obs e Itens ao Pedido Anterior]
    C -->|Não| E[Extrai Bounding Boxes]
    E --> F[Fallback via Regex (Endereços)]
    F --> G[Normaliza Locais e ID do Cliente]
    G --> H[Objeto Python Dicionário]"""
    add_mermaid_image(doc, pdf_arch_code, 'pdf_arch.png')
    doc.add_paragraph('Uso estratégico da biblioteca pdfplumber aplicando geometria de Bounding Boxes combinada com parsing de Regex, capturando anomalias em faturas compostas por múltiplas páginas.')

    # Módulos e Endpoints
    add_heading(doc, '2.2 Módulos Internos (Endpoints Lógicos)', 2)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('excel_handler.py: ').bold = True
    p.add_run('Coração do backend. Responsável pelas operações CRUD. Funções chaves: `read_principal()` (leitura em dataframe pandas), `update_status_batch()` (alteração em massa de status), e `archive_completed()` (migração de Dados para aba de Histórico).')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('pdf_parser.py: ').bold = True
    p.add_run('Processamento visual do PDF. Define variáveis globais como `BBOX_PEDIDO` e `BBOX_CLIENTE`. Possui a função `_is_continuation_page()` vital para lidar com pedidos imensos da operação que transbordam a folha.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('expedition_engine.py / checklist_engine.py: ').bold = True
    p.add_run('Módulos de negócio acionados nas telas 5 e 8, processando lógicas específicas como "associação de veículos a NFs" sem sujar a lógica de apresentação (app.py).')

    # Dependências
    add_heading(doc, '2.3 Dependências (requirements.txt)', 2)
    doc.add_paragraph('As bibliotecas principais do ecossistema são:')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('streamlit (>=1.57.0): ').bold = True
    p.add_run('Framework principal do frontend web.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('openpyxl / pandas: ').bold = True
    p.add_run('I/O de arquivos Microsoft Excel de forma programática (DataFrames e Células).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('pdfplumber: ').bold = True
    p.add_run('Extração profunda de texto baseado em localização espacial nos PDFs.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('folium / geopy: ').bold = True
    p.add_run('Mapeamento espacial, transformando texto de logradouro em LAT/LONG e gerando Leaflet Maps.')

    # Considerações Finais
    add_heading(doc, 'Considerações Finais (Refinamento)', 1, RGBColor(33, 150, 243))
    doc.add_paragraph('A divisão semântica deste software prova que é possível aliar alta resiliência de banco de dados nativos da indústria (Excel) à velocidade e acessibilidade de micro-frameworks (Streamlit) para criar uma operação logística escalável e acessível.')

    doc.save('Documentacao_Controle_Notas.docx')
    print('Documento DOCX reestruturado gerado com sucesso!')

if __name__ == '__main__':
    create_doc()
